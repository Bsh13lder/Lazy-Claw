"""Tests for rich docx export/import: lists, headings, bold (Task A4)."""

from __future__ import annotations

import io

from docx import Document

from lazyclaw.docs import snapshot as D
from lazyclaw.docs.docx_io import snapshot_to_docx, docx_to_snapshot


def _snap(blocks):
    return {"id": "doc-1", "documentStyle": {}, "body": D.build_body_with_blocks(blocks)}


def test_numbered_list_exports_as_list_number_style():
    snap = _snap(
        [
            {"type": "number", "level": 0, "runs": [{"text": "first"}]},
            {"type": "number", "level": 0, "runs": [{"text": "second"}]},
        ]
    )
    doc = Document(io.BytesIO(snapshot_to_docx(snap)))
    styles = [p.style.name for p in doc.paragraphs if p.text]
    assert styles and all("List Number" in s for s in styles)


def test_bullet_list_exports_as_list_bullet_style():
    snap = _snap([{"type": "bullet", "level": 0, "runs": [{"text": "x"}]}])
    doc = Document(io.BytesIO(snapshot_to_docx(snap)))
    assert any("List Bullet" in p.style.name for p in doc.paragraphs if p.text)


def test_heading_exports_as_heading_style():
    snap = _snap([{"type": "heading", "level": 1, "runs": [{"text": "Title"}]}])
    doc = Document(io.BytesIO(snapshot_to_docx(snap)))
    assert any("Heading 1" in p.style.name for p in doc.paragraphs if p.text)


def test_bold_run_exports_bold():
    snap = _snap(
        [{"type": "paragraph", "level": 0, "runs": [{"text": "x", "bold": True}]}]
    )
    doc = Document(io.BytesIO(snapshot_to_docx(snap)))
    assert any(r.bold for p in doc.paragraphs for r in p.runs)


def test_docx_list_roundtrips_to_number_block():
    snap = _snap([{"type": "number", "level": 0, "runs": [{"text": "a"}]}])
    back = docx_to_snapshot(snapshot_to_docx(snap), "RT")
    assert D.get_blocks(back)[0]["type"] == "number"


def test_docx_heading_roundtrips_to_heading_block():
    snap = _snap([{"type": "heading", "level": 2, "runs": [{"text": "Sec"}]}])
    back = docx_to_snapshot(snapshot_to_docx(snap), "RT")
    blocks = D.get_blocks(back)
    assert blocks[0]["type"] == "heading" and blocks[0]["level"] == 2


def test_docx_bold_roundtrips():
    snap = _snap(
        [{"type": "paragraph", "level": 0, "runs": [{"text": "strong", "bold": True}]}]
    )
    back = docx_to_snapshot(snapshot_to_docx(snap), "RT")
    runs = D.get_blocks(back)[0]["runs"]
    assert any(r.get("bold") for r in runs)


def test_link_still_roundtrips():
    snap = _snap(
        [{"type": "paragraph", "level": 0,
          "runs": [{"text": "see ", }, {"text": "site", "url": "https://x.io"}]}]
    )
    back = docx_to_snapshot(snapshot_to_docx(snap), "RT")
    runs = D.get_blocks(back)[0]["runs"]
    assert any(r.get("url") == "https://x.io" for r in runs)
