"""Hyperlink round-trip through the .docx layer (lazyclaw/docs/docx_io.py).

python-docx has no native hyperlink API, so we emit/parse ``w:hyperlink``
elements ourselves. These tests prove a hyperlink survives snapshot → .docx →
snapshot with its label and URL intact, and that the exported .docx carries a
real external relationship.
"""

from __future__ import annotations

import io

from docx import Document

from lazyclaw.docs import docx_io, snapshot as D


def _snap_with_link():
    body = D.build_body_with_runs(
        [
            [{"text": "Heading line"}],
            [{"text": "Visit "}, {"text": "my site", "url": "https://example.com"}, {"text": " today"}],
        ]
    )
    return {"id": "doc-x", "name": "Linked", "documentStyle": {}, "body": body}


def test_export_emits_real_hyperlink_relationship():
    data = docx_io.snapshot_to_docx(_snap_with_link())
    doc = Document(io.BytesIO(data))
    # the body paragraph should carry a w:hyperlink element
    xml = "".join(p._p.xml for p in doc.paragraphs)
    assert "hyperlink" in xml
    # and an external relationship to the url must exist somewhere in the part
    targets = [r.target_ref for r in doc.part.rels.values()]
    assert "https://example.com" in targets


def test_export_then_import_round_trips_link():
    data = docx_io.snapshot_to_docx(_snap_with_link())
    snap2 = docx_io.docx_to_snapshot(data, name="Linked")
    runs = D.get_paragraph_runs(snap2)
    # find the link run
    flat = [r for para in runs for r in para]
    link = [r for r in flat if r.get("url")]
    assert link == [{"text": "my site", "url": "https://example.com"}]
    # plain text preserved around it
    assert "Visit my site today" in D.get_text(snap2)


def test_plain_doc_still_round_trips():
    base = D.set_text(D.blank_document("Plain"), "one\ntwo\nthree")
    data = docx_io.snapshot_to_docx(base)
    out = docx_io.docx_to_snapshot(data, name="Plain")
    assert D.get_paragraphs(out) == ["one", "two", "three"]
    # no spurious customRanges on a link-free doc
    assert out["body"].get("customRanges", []) == []


def test_import_docx_without_links_has_no_customranges():
    document = Document()
    document.add_paragraph("just plain text")
    buf = io.BytesIO()
    document.save(buf)
    out = docx_io.docx_to_snapshot(buf.getvalue(), name="P")
    assert "just plain text" in D.get_text(out)
    assert out["body"].get("customRanges", []) == []
