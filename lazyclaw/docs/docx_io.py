"""Convert between the Univer ``IDocumentData`` snapshot and ``.docx`` / PDF.

The snapshot is LazyClaw's source of truth; ``.docx`` is a derived format
produced on export and parsed on import. PDF is a best-effort secondary export
that shells out to LibreOffice headless when it is available on the host.

python-docx (MIT) handles the ``.docx`` direction at the paragraph level — we
walk :func:`lazyclaw.docs.snapshot.get_paragraphs` on the way out and feed the
read-back paragraphs through :func:`lazyclaw.docs.snapshot.set_text` on the way
in, so the plain-text model round-trips cleanly.
"""

from __future__ import annotations

import io
import logging
import os
import shutil
import subprocess  # nosec B404 — fixed args, no shell, used for LibreOffice convert
import tempfile
from typing import Any

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from lazyclaw.docs import snapshot as D

logger = logging.getLogger(__name__)

# How long to give LibreOffice before giving up on a PDF conversion.
_SOFFICE_TIMEOUT_S = 60


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append a real ``w:hyperlink`` run to a python-docx paragraph.

    python-docx has no hyperlink API, so we build the OOXML element directly:
    relate the paragraph's part to the external ``url`` and wrap a run in a
    ``w:hyperlink`` carrying that relationship id. Styled blue + underlined so
    it reads as a link.
    """
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


# Map a block type → python-docx paragraph style (present in the default
# template). Headings use "Heading {level}".
_LIST_STYLE = {"bullet": "List Bullet", "number": "List Number"}


def snapshot_to_docx(snap: dict[str, Any]) -> bytes:
    """Render a Univer document snapshot to ``.docx`` bytes with real formatting.

    Walks the typed blocks (:func:`lazyclaw.docs.snapshot.get_blocks`): headings
    use the ``Heading N`` style, ordered/unordered lists use ``List Number`` /
    ``List Bullet``, and runs carry bold/italic/underline. Runs with a ``url``
    become real ``w:hyperlink`` elements. An empty document still produces a
    valid ``.docx``.
    """
    document = Document()
    for block in D.get_blocks(snap):
        btype = block.get("type", "paragraph")
        runs = block.get("runs") or []
        if btype == "heading":
            level = max(1, min(3, int(block.get("level") or 1)))
            para = document.add_paragraph(style=f"Heading {level}")
        elif btype in _LIST_STYLE:
            para = document.add_paragraph(style=_LIST_STYLE[btype])
        else:
            para = document.add_paragraph()
        for run in runs:
            url = run.get("url")
            if url:
                _add_hyperlink(para, run.get("text", ""), str(url))
            else:
                r = para.add_run(run.get("text", ""))
                if run.get("bold"):
                    r.bold = True
                if run.get("italic"):
                    r.italic = True
                if run.get("underline"):
                    r.underline = True

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _run_flags_from_rpr(rpr) -> dict[str, bool]:
    """Read bold/italic/underline flags from a ``w:r``'s ``w:rPr`` element."""
    flags: dict[str, bool] = {}
    if rpr is None:
        return flags

    def _on(tag: str) -> bool:
        el = rpr.find(qn(tag))
        if el is None:
            return False
        val = el.get(qn("w:val"))
        return val not in ("false", "0", "none")

    if _on("w:b"):
        flags["bold"] = True
    if _on("w:i"):
        flags["italic"] = True
    if _on("w:u"):
        flags["underline"] = True
    return flags


def _docx_paragraph_runs(paragraph) -> list[dict[str, Any]]:
    """Extract ordered runs (plain + hyperlink, with emphasis) from a paragraph."""
    runs: list[dict[str, Any]] = []
    rels = paragraph.part.rels
    for child in paragraph._p:
        tag = child.tag
        if tag == qn("w:r"):
            text = "".join(node.text or "" for node in child.findall(qn("w:t")))
            if text:
                run: dict[str, Any] = {"text": text}
                run.update(_run_flags_from_rpr(child.find(qn("w:rPr"))))
                runs.append(run)
        elif tag == qn("w:hyperlink"):
            r_id = child.get(qn("r:id"))
            url = None
            if r_id and r_id in rels:
                url = rels[r_id].target_ref
            text = "".join(node.text or "" for node in child.iter(qn("w:t")))
            if text:
                runs.append({"text": text, "url": url} if url else {"text": text})
    return runs


def _block_type_from_style(style_name: str) -> tuple[str, int]:
    """Map a python-docx paragraph style name → (block type, level)."""
    name = style_name or ""
    if name.startswith("List Number"):
        return "number", 0
    if name.startswith("List Bullet"):
        return "bullet", 0
    if name.startswith("Heading"):
        digits = "".join(ch for ch in name if ch.isdigit())
        level = int(digits) if digits else 1
        return "heading", max(1, min(3, level))
    return "paragraph", 0


def docx_to_snapshot(data: bytes, name: str | None = None) -> dict[str, Any]:
    """Parse ``.docx`` bytes into a Univer snapshot (lists/headings/emphasis)."""
    document = Document(io.BytesIO(data))
    blocks: list[dict[str, Any]] = []
    for p in document.paragraphs:
        runs = _docx_paragraph_runs(p)
        try:
            style_name = p.style.name if p.style else ""
        except Exception:  # noqa: BLE001 — defensive against odd templates
            style_name = ""
        btype, level = _block_type_from_style(style_name)
        blocks.append({"type": btype, "level": level, "runs": runs})
    # Drop trailing empty paragraphs python-docx often emits.
    while blocks and not blocks[-1]["runs"]:
        blocks.pop()
    base = D.blank_document(name or "Imported")
    if not blocks:
        return base
    return {**base, "body": D.build_body_with_blocks(blocks)}


# ───────────────────────── PDF (best-effort) ────────────────────────

def _libreoffice_available() -> bool:
    """Return True when a LibreOffice headless binary is on the PATH."""
    return any(shutil.which(name) for name in ("soffice", "libreoffice"))


def _soffice_binary() -> str | None:
    for name in ("soffice", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path
    return None


def snapshot_to_pdf(snap: dict[str, Any]) -> bytes | None:
    """Best-effort PDF export via LibreOffice headless.

    Writes a temp ``.docx``, runs ``soffice --headless --convert-to pdf`` on
    it, and returns the resulting PDF bytes. Returns ``None`` (never raises)
    when LibreOffice is absent or the conversion fails for any reason — callers
    should treat ``None`` as "PDF unavailable on this host".
    """
    binary = _soffice_binary()
    if not binary:
        logger.debug("snapshot_to_pdf skipped: no LibreOffice binary on PATH")
        return None

    docx_bytes = snapshot_to_docx(snap)
    try:
        with tempfile.TemporaryDirectory() as workdir:
            docx_path = os.path.join(workdir, "doc.docx")
            with open(docx_path, "wb") as fh:
                fh.write(docx_bytes)
            result = subprocess.run(  # nosec B603 — fixed binary + args, no shell
                [
                    binary,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    workdir,
                    docx_path,
                ],
                cwd=workdir,
                capture_output=True,
                timeout=_SOFFICE_TIMEOUT_S,
                check=False,
            )
            pdf_path = os.path.join(workdir, "doc.pdf")
            if result.returncode != 0 or not os.path.exists(pdf_path):
                logger.warning(
                    "LibreOffice PDF conversion failed (rc=%s): %.200s",
                    result.returncode,
                    result.stderr.decode("utf-8", "replace") if result.stderr else "",
                )
                return None
            with open(pdf_path, "rb") as fh:
                return fh.read()
    except Exception as exc:  # noqa: BLE001 — best-effort, never propagate
        logger.warning("snapshot_to_pdf failed: %s", exc)
        return None
